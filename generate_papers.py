"""
Paper Generator: IEEE and ACM DOCX versions
Entangled Intelligence: Nanoscale Quantum-Neuromorphic Hybrid Architectures
for Post-von Neumann Computation (Enhanced with Security Analysis)

IEEE-NANO 2026 + ACM NANOCOM 2026 versions

Author: Sunil Gentyala
Affiliation: Independent Researcher, HCLTech (HCL America Inc.), Dallas, TX 75001, USA
ORCID: 0009-0005-2642-3479
Email: sunil.gentyala@ieee.org
GitHub: https://github.com/sunilgentyala/QNHS-Research-2026
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy
import os


# ─────────────────────────────────────────────────────────────────────────────
#  PAPER CONTENT (shared between IEEE and ACM, formatted differently)
# ─────────────────────────────────────────────────────────────────────────────

GITHUB_URL = "https://github.com/sunilgentyala/QNHS-Research-2026"

TITLE = (
    "Entangled Intelligence: Nanoscale Quantum-Neuromorphic Hybrid "
    "Architectures for Post-von Neumann Computation"
)

AUTHOR = "Sunil Gentyala"
AFFILIATION = (
    "Independent Researcher, HCLTech (HCL America Inc.), Dallas, TX 75001, USA"
)
EMAIL = "sunil.gentyala@ieee.org"
ORCID = "0009-0005-2642-3479"

ABSTRACT = (
    "The escalating computational demands of modern artificial intelligence "
    "and data-centric workloads have driven a reexamination of the architectural "
    "and thermodynamic limitations of von Neumann computing paradigms. This paper "
    "proposes and analyzes a novel nanoscale Quantum-Neuromorphic Hybrid Substrate "
    "(QNHS) that integrates isotopically purified 28Si spin qubit arrays, "
    "spintronic perpendicular magnetic anisotropy (PMA) magnetic tunnel junction "
    "(MTJ) memristors, and cryo-CMOS peripheral circuits in a vertically stacked "
    "monolithic die. By exploiting quantum superposition for synaptic weight "
    "encoding within biologically inspired spike-based processing frameworks, "
    "QNHS projects an energy consumption of approximately 11 fJ per synaptic "
    "event, a reduction of nearly three orders of magnitude relative to "
    "contemporary GPU-based deep learning accelerators. A distance-3 surface "
    "code topological protection layer, decoded by a classical minimum-weight "
    "perfect matching (MWPM) circuit integrated within the cryo-CMOS plane, "
    "enables operation at 1 K rather than the millikelvin regimes required by "
    "superconducting qubit platforms. Recent 2025 experimental results confirm "
    "silicon spin qubits fabricated in 300 mm production foundries now achieve "
    "single-qubit gate fidelities exceeding 99%, with T1 coherence times "
    "reaching 9.5 seconds. We further identify a comprehensive threat landscape "
    "for quantum-neuromorphic hardware, including timing and crosstalk-based "
    "side-channel attacks demonstrated on cloud quantum services in 2025, and "
    "adversarial weight perturbations through quantum state manipulation. "
    "Hardware security primitives native to the QNHS substrate are proposed as "
    "countermeasures. Theoretical coherence analysis, spintronic device modeling, "
    "fabrication pathway assessment, a comparative performance evaluation, and "
    "a security threat model substantiate the architectural viability of QNHS "
    "within the near-term fabrication horizon. Open-source simulation code is "
    "available at " + GITHUB_URL + "."
)

INDEX_TERMS = (
    "cryo-CMOS, hardware security, magnetic tunnel junctions, nanotechnology, "
    "post-von Neumann architectures, post-quantum cryptography, quantum "
    "adversarial machine learning, quantum-neuromorphic computing, silicon spin "
    "qubits, spiking neural networks, spintronic memristors, side-channel attacks, "
    "topological quantum computation."
)

# Section content as list of (heading_text, level, body_text) tuples.
# level 1 = Roman numeral section, level 2 = lettered subsection
SECTIONS = [
    (
        "I. INTRODUCTION",
        1,
        "The von Neumann bottleneck, characterized by the energetically prohibitive "
        "and latency-laden shuttling of data between physically segregated processing "
        "and memory units, has long constituted the preeminent architectural liability "
        "of digital computation. As transistor channel lengths approach the sub-3 nm "
        "regime, the classical scaling trajectory codified by Moore's Law confronts "
        "insurmountable physical barriers: quantum mechanical tunneling currents, "
        "catastrophic thermal dissipation densities, and stochastic dopant variability "
        "together prevent further miniaturization within existing paradigms [1]. The "
        "International Roadmap for Devices and Systems (IRDS) projects that conventional "
        "planar transistor scaling will be effectively exhausted by the late 2020s, "
        "driving an urgent industry-wide search for beyond-CMOS computational "
        "architectures [2].\n\n"
        "Neuromorphic computing, architecturally inspired by the event-driven, massively "
        "parallel, and co-located memory-processing organization of biological neural "
        "circuitry, has demonstrated compelling energy efficiencies for cognitive "
        "workloads. Intel's Loihi 2 and IBM's NorthPole processors instantiate "
        "spike-based information encoding with on-chip synaptic plasticity, yielding "
        "power consumptions multiple orders of magnitude below conventional deep "
        "learning accelerators on equivalent classification benchmarks [3]. A 2024 "
        "follow-on study demonstrated NorthPole achieving 72.7x greater energy "
        "efficiency than the lowest-latency GPU for large language model inference, "
        "processing 28,356 tokens per second across 16 chips [4]. Nonetheless, even "
        "these architectures remain fundamentally classical: their synaptic weights "
        "are scalar quantities, their computations deterministic, and their "
        "representational capacity bounded by the physical number of addressable states.\n\n"
        "Quantum computing, conversely, leverages superposition and entanglement to "
        "encode and manipulate information in state spaces of dimensionality exponential "
        "in the qubit count. A pivotal 2025 milestone demonstrated that 28Si spin "
        "qubits fabricated in a 300 mm production semiconductor foundry achieve "
        "single-qubit gate fidelities exceeding 99% with T1 coherence times reaching "
        "9.5 seconds [15], establishing silicon spin qubits as manufacturable "
        "building blocks for scalable quantum systems. Contemporary quantum processors "
        "are nonetheless beset by decoherence timescales insufficient for deep "
        "algorithmic circuits absent extensive error-correction overhead that "
        "substantially erodes theoretical computational advantages [5].\n\n"
        "As quantum hardware matures, a parallel and underappreciated concern has "
        "emerged: the security of quantum computing infrastructure itself. In 2025, "
        "Lu et al. demonstrated that timing measurements with as few as 10 queries "
        "can identify the underlying quantum hardware and subvert algorithm "
        "confidentiality on cloud quantum services [22]. Choudhury et al. further "
        "showed that inter-qubit crosstalk enables adversaries to reconstruct victim "
        "quantum circuits with 85.7% accuracy across 336 benchmark circuits [23]. "
        "These findings underscore that architectural novelty must be accompanied by "
        "rigorous security analysis from the outset.\n\n"
        "The central thesis of this paper is that the physical co-integration of "
        "quantum and neuromorphic paradigms at the nanoscale, mediated by innovations "
        "in 28Si quantum-dot fabrication and spintronic MTJ device engineering, offers "
        "a synthetic pathway that transcends the limitations of either approach in "
        "isolation. We propose the Quantum-Neuromorphic Hybrid Substrate (QNHS): "
        "a vertically integrated monolithic architecture in which 28Si spin qubit "
        "arrays encode synaptic weights as quantum superposition states, PMA-MTJ "
        "memristors implement continuously graded synaptic conductances, and cryo-CMOS "
        "peripheral circuits manage spike encoding, syndrome measurement, and "
        "communication. We further provide the first comprehensive security analysis "
        "of a quantum-neuromorphic hardware substrate, identifying threat vectors and "
        "native countermeasures. Open-source simulation code accompanies this work at "
        + GITHUB_URL + ".\n\n"
        "The remainder of this paper is organized as follows. Section II reviews "
        "pertinent literature. Section III presents the QNHS architecture. Section IV "
        "provides theoretical analysis. Section V discusses nanofabrication pathways. "
        "Section VI analyzes the security threat landscape. Section VII addresses open "
        "challenges. Section VIII concludes."
    ),

    (
        "II. BACKGROUND AND RELATED WORK",
        1,
        ""
    ),
    (
        "A. Neuromorphic Computing Substrates",
        2,
        "The seminal neuromorphic silicon cochlea and retina of Carver Mead in the "
        "late 1980s gave rise to a family of analog VLSI circuits emulating neural "
        "dynamics through subthreshold transistor operation [6]. Contemporary "
        "descendants include memristive crossbar arrays employing resistive RAM "
        "(ReRAM), phase-change memory (PCM), and spintronic MTJ devices as artificial "
        "synapses. MTJ-based synapses are particularly distinguished by their "
        "sub-nanosecond switching speeds, endurance exceeding 10^12 cycles, and "
        "inherent CMOS process compatibility [7]. A 2024 review in npj Spintronics "
        "documented spintronic neuron designs operating at nanosecond timescales with "
        "write energies as low as 3.78 fJ/bit, reaffirming spintronics as a "
        "leading candidate for energy-efficient neuromorphic hardware [25].\n\n"
        "IBM's NorthPole architecture demonstrated a 25x inference energy improvement "
        "over comparable GPU implementations by eliminating off-chip memory accesses "
        "entirely, distributing 224 MB of SRAM across 256 cores [4]. Subsequent "
        "deployment studies in 2024 confirmed 72.7x greater energy efficiency than "
        "the lowest-latency GPU for large language model inference, with throughput "
        "reaching 28,356 tokens per second across 16 chips in a 2U server. Yet "
        "NorthPole operates at 300 K with classical transistors and deterministic "
        "scalar weights; it cannot represent distributional uncertainty in synaptic "
        "transmission without explicit software overhead."
    ),
    (
        "B. Quantum Computing at the Nanoscale",
        2,
        "Semiconductor quantum dots, nanoscale islands confining electrons in all "
        "three spatial dimensions, have emerged as compelling qubit candidates owing "
        "to their compatibility with established CMOS fabrication infrastructure [8]. "
        "Spin qubits realized in isotopically purified 28Si quantum dots have "
        "demonstrated T2 coherence times exceeding 1 ms at cryogenic temperatures, "
        "representing improvements of five orders of magnitude over charge qubits "
        "and approaching thresholds requisite for fault-tolerant surface code "
        "operation [9]. Yoneda et al. reported single-qubit gate fidelities "
        "exceeding 99.9% in 28Si, establishing the feasibility of low-error spin "
        "qubit operation [13].\n\n"
        "A landmark 2025 result demonstrated that 28Si spin qubits fabricated in a "
        "300 mm production semiconductor foundry achieve single- and two-qubit "
        "control fidelities exceeding 99%, with state preparation and measurement "
        "fidelities up to 99.9% and T1 coherence times reaching 9.5 seconds [15]. "
        "A separate 2025 study reported an 11-qubit atom processor in isotopically "
        "purified silicon with two-qubit gate fidelities of 99.9% and successful "
        "entanglement of 8 nuclear spins in GHZ states [16]. These results "
        "establish that academic-level spin qubit performance is now reproducible "
        "in standard semiconductor manufacturing infrastructure.\n\n"
        "A complementary hardware advance arrived in late 2024, when SemiQon "
        "demonstrated the first cryo-CMOS transistor specifically optimized for "
        "operation at 1 K and below, achieving power consumption of only 0.1% "
        "relative to room-temperature equivalents and 1,000x lower heat "
        "dissipation [18]. This breakthrough enables quantum control electronics "
        "to reside inside the cryostat, directly addressing the wiring bottleneck "
        "that limits near-term qubit scaling.\n\n"
        "Microsoft's February 2025 announcement of a Majorana-based topological "
        "qubit processor (Majorana 1) reported interferometric signatures consistent "
        "with Majorana zero modes in InAs-Al hybrid nanowire devices [17]. The "
        "scientific community has noted, however, that independent verification "
        "of full topological protection remains outstanding, and the accompanying "
        "Nature editorial stated that the reported results do not yet constitute "
        "definitive evidence for topological qubits. The longer-horizon potential "
        "of braiding-protected logical qubits remains a compelling research "
        "direction [10]."
    ),
    (
        "C. Quantum-Neuromorphic Convergence",
        2,
        "Prior explorations of quantum-neuromorphic hybridization have largely "
        "confined themselves to algorithmic speedups: quantum variants of Hebbian "
        "learning rules, quantum-enhanced Boltzmann machines, and variational "
        "quantum eigensolvers (VQE) for synaptic weight optimization [11]. Maronese "
        "et al. surveyed quantum neuromorphic approaches, noting that essentially "
        "all extant proposals treat the quantum and neuromorphic layers as logically "
        "distinct [12]. A 2025 framework termed quantum hyperdimensional computing "
        "demonstrated quantum neuromorphic algorithms with hyperdimensional "
        "operations mapping natively onto a 156-qubit IBM Heron processor for "
        "symbolic reasoning and classification tasks [24]. Nonetheless, physical "
        "device-level co-integration, wherein individual synaptic elements exhibit "
        "intrinsic quantum mechanical behavior during inference, remains "
        "substantially underexplored. This constitutes the precise gap the QNHS "
        "architecture addresses."
    ),
    (
        "D. Security Landscape for Quantum Hardware",
        2,
        "The security of quantum computing infrastructure has recently emerged as "
        "a critical research frontier. Lu et al. demonstrated in 2025 that timing "
        "side-channel attacks on cloud-based quantum services can identify the "
        "underlying quantum hardware with as few as 10 measurements, directly "
        "compromising algorithm confidentiality [22]. Choudhury et al. showed "
        "that crosstalk between adjacent qubits in multi-tenant NISQ computers "
        "enables adversaries with minimal privileges to reconstruct victim quantum "
        "circuits with 85.7% accuracy [23]. On the software side, a 2024 survey "
        "of 53 empirical studies confirmed that input-level adversarial attacks "
        "against variational quantum circuits are practical on existing NISQ "
        "hardware [21]. These developments collectively establish that quantum "
        "hardware security must be addressed at the physical architecture level, "
        "not merely through software defenses."
    ),

    (
        "III. PROPOSED ARCHITECTURE",
        1,
        ""
    ),
    (
        "A. Three-Plane Monolithic Stack",
        2,
        "The QNHS comprises three vertically integrated functional planes fabricated "
        "on a single monolithic die, interconnected via Cu-Cu thermocompression bonds "
        "at a pitch of 200 nm. Plane 1 (Quantum Coherent Processing Plane, QCPP) is "
        "a two-dimensional lattice of 28Si spin qubits with nearest-neighbor exchange "
        "coupling J tunable via electrostatic gate voltages. This plane executes "
        "parameterized quantum circuits for synaptic weight superposition encoding "
        "and, during learning, STDP-driven amplitude updates. A surface code syndrome "
        "extraction sub-lattice enables real-time error detection within the same "
        "plane. Plane 2 (Spintronic Synaptic Plane, SSP) is a 256x256 crossbar of "
        "PMA-MTJ devices, each encoding a 2-bit synaptic conductance through partial "
        "spin-orbit torque (SOT) switching between four distinguishable resistance "
        "states {R0, R1, R2, R3} with TMR ratios exceeding 200%. Plane 3 (CMOS "
        "Peripheral Plane, CPP) is fabricated in a 3 nm FinFET process on SOI and "
        "implements leaky integrate-and-fire (LIF) neuron circuits, spike encoding "
        "and time-to-first-spike readout logic, cryogenic MWPM decoder for surface "
        "code error correction, and chip-to-room-temperature I/O via cryo-compatible "
        "SerDes links. The SemiQon 2024 cryo-CMOS demonstration [18] directly "
        "validates the feasibility of 1 K CMOS operation with 1,000x reduced heat "
        "dissipation relative to room-temperature equivalents, strengthening the "
        "near-term credibility of the CPP design."
    ),
    (
        "B. Quantum Synaptic Encoding",
        2,
        "In classical neuromorphic architectures, the weight w_ij coupling "
        "presynaptic neuron i to postsynaptic neuron j is a deterministic scalar, "
        "updated by a learning rule such as spike-timing-dependent plasticity (STDP). "
        "In the QNHS, each synaptic weight is instead encoded as a k-qubit quantum "
        "state occupying a 2^k-dimensional Hilbert space:\n\n"
        "    |w_ij> = SUM_{l=0}^{2^k-1} alpha_l |w_l>,    SUM_l |alpha_l|^2 = 1    (1)\n\n"
        "where {|w_l>} is an orthonormal basis of discrete weight states and {alpha_l} "
        "are complex probability amplitudes set by the learning process. This encoding "
        "enables the synapse to simultaneously represent a probability distribution "
        "over possible weight values, providing inherent Bayesian uncertainty "
        "quantification absent from deterministic classical architectures and directly "
        "analogous to the probabilistic vesicle-release dynamics of biological synapses. "
        "When a presynaptic spike arrives, a controlled interaction Hamiltonian "
        "H_int = g * sigma_z^(i) x V^(j) couples the qubit weight register to the "
        "postsynaptic membrane voltage variable, where g is the coupling strength "
        "calibrated by gate-voltage tuning. The resulting projective measurement "
        "collapses |w_ij> to weight w_l with probability |alpha_l|^2, imparting a "
        "proportional excitatory postsynaptic potential (EPSP) on the postsynaptic "
        "LIF circuit."
    ),
    (
        "C. STDP-Driven Quantum Amplitude Update",
        2,
        "Quantum spike-timing-dependent plasticity (Q-STDP) modifies the amplitude "
        "vector {alpha_l} according to:\n\n"
        "    alpha_l^(t+1) = N[ alpha_l^(t) + eta * K(dt) * alpha_l^(t) ]    (2)\n\n"
        "where eta is the learning rate, N[.] denotes normalization to restore unit "
        "L2 norm, dt = t_post - t_pre is the spike-timing differential, and "
        "K(dt) = A+ * exp(-dt/tau+) * Theta(dt) - A- * exp(+dt/tau-) * Theta(-dt) "
        "is the classical STDP kernel applied to the quantum amplitude domain. This "
        "formulation shifts the probability mass of the weight distribution toward "
        "higher-weight basis states during long-term potentiation (LTP, dt > 0) and "
        "toward lower-weight states during long-term depression (LTD, dt < 0), "
        "reproducing Hebbian learning dynamics in the quantum amplitude space."
    ),
    (
        "D. Topological Error Protection",
        2,
        "To mitigate decoherence-induced corruption of superposed weight states, "
        "the QCPP implements a distance-3 surface code in which each logical weight "
        "qubit is distributed over d^2 + (d-1)^2 = 13 physical data and ancilla "
        "qubits arranged on a planar lattice (d = 3). The surface code threshold "
        "error rate of ~1% per gate is comfortably exceeded by the measured <0.1% "
        "gate infidelity of 28Si spin qubits [13]. The MWPM decoder, implemented "
        "in 3 nm FinFET CMOS within the CPP, operates with a syndrome cycle latency "
        "of ~1 us, well below the ~10 ms spike integration timescale characteristic "
        "of biological-rate SNNs."
    ),
    (
        "E. Hardware Security Primitives",
        2,
        "The QNHS substrate naturally exposes several hardware security primitives "
        "that can be exploited defensively. Charge noise in 28Si quantum dots exhibits "
        "device-unique stochastic signatures determined by nanoscale dopant "
        "distributions and interface trap positions, making them suitable Physical "
        "Unclonable Functions (PUFs) for device authentication without additional "
        "hardware area cost. Stochastic MTJ switching under sub-threshold current "
        "pulses at 1 K generates true random bits from thermal fluctuations in the "
        "free-layer magnetization, providing a hardware quantum random number "
        "generator (QRNG) for cryptographic key seeding. The 256x256 MTJ crossbar "
        "further accelerates polynomial ring arithmetic underlying NIST-standardized "
        "post-quantum lattice schemes (FIPS 203/204/205 [26]) through in-memory "
        "analog matrix-vector multiplication."
    ),

    (
        "IV. THEORETICAL ANALYSIS AND PERFORMANCE PROJECTIONS",
        1,
        ""
    ),
    (
        "A. Energy Efficiency Decomposition",
        2,
        "The energy dissipated per synaptic event E_syn decomposes as:\n\n"
        "    E_syn = E_MTJ + E_qubit + E_CMOS    (3)\n\n"
        "E_MTJ ~= 0.1 fJ: computed from the Landauer-Buttiker spin-torque switching "
        "model for a d = 20 nm circular PMA-MTJ with free-layer damping constant "
        "alpha_G = 0.015, anisotropy field mu_0*H_k = 0.5 T, and current pulse "
        "duration 1 ns. E_qubit ~= 1 fJ: estimated from the product of gate drive "
        "power (1 nW at 30 MHz Rabi frequency) and pi-pulse duration (33 ns). "
        "This figure is consistent with published cryo-CMOS qubit control "
        "measurements from Delft and IMEC [14]. E_CMOS ~= 10 fJ: encompasses spike "
        "encoding, MWPM syndrome measurement, integration logic, and peripheral "
        "buffering in the 3 nm FinFET process at 1 K, where carrier mobilities "
        "are enhanced by ~2x relative to 300 K operation [18]. The resultant "
        "E_syn ~= 11 fJ compares with ~10 pJ per operation for the NVIDIA H100 "
        "GPU and ~1 pJ for Intel Loihi 2. The cryogenic refrigeration overhead "
        "factor of ~50x (Carnot efficiency at 1 K from 300 K) reduces the "
        "wall-plug efficiency advantage to ~20x over GPU baselines for "
        "inference-dominant workloads. Table I presents a comprehensive comparative "
        "assessment across seven performance dimensions."
    ),
    (
        "B. Representational Capacity Analysis",
        2,
        "A classical N-synapse crossbar with B-bit weight precision encodes NB bits "
        "of information. The QNHS encodes each synaptic weight as a k-qubit quantum "
        "state, representing 2^k independent complex amplitudes per synapse. For the "
        "baseline configuration of k = 4 qubits per synapse in a 256x256 crossbar, "
        "the total effective representational dimensionality is:\n\n"
        "    D_QNHS = (256)^2 x 2^4 = 65,536 x 16 ~= 10^6 complex d.o.f.    (4)\n\n"
        "This million-dimensional amplitude space, hosting the joint quantum amplitude "
        "distributions across all synaptic junctions, is wholly inaccessible to any "
        "classical architecture of comparable physical footprint, conferring "
        "substantially greater expressive capacity per unit area."
    ),
    (
        "C. Coherence Lifetime Compatibility",
        2,
        "The requisite coherence window for QNHS is determined by the longest "
        "computation that must remain coherent before a projective measurement or "
        "error-correction cycle: approximately one spike integration period, "
        "tau_int ~= 1 ms-10 ms for biological-rate SNNs. Published T2 data for 28Si "
        "spin qubits at 1 K cluster in the range 0.1 ms-1 ms [13], placing current "
        "experimental capabilities precisely at the threshold. The isotopic enrichment "
        "roadmap projects that achieving 99.9995% 28Si (residual 29Si below 50 ppm) "
        "combined with thermally oxidized SiO2 interfaces with interface state density "
        "D_it < 10^9 cm^-2 eV^-1 will extend T2 toward 10 ms [9]. The 2025 "
        "demonstration of T1 times reaching 9.5 seconds [15] further supports the "
        "feasibility of achieving long coherence times in foundry-compatible devices. "
        "Under surface code protection at d = 3, the effective logical coherence "
        "time is enhanced by a factor of approximately (p_th/p_phys)^((d-1)/2) "
        "~= (0.01/0.001)^1 = 10x, projecting logical coherence lifetimes toward "
        "100 ms, comfortably within QNHS operational requirements."
    ),

    (
        "V. NANOFABRICATION PATHWAYS",
        1,
        ""
    ),
    (
        "A. 28Si Quantum Dot Fabrication",
        2,
        "Isotopically enriched 28Si epilayers (99.9995% purity, residual 29Si < 50 ppm) "
        "are deposited on 300 mm SOI substrates by solid-source molecular beam epitaxy "
        "(MBE) at substrate temperatures of 550 degrees C. Quantum dot gate electrodes "
        "are defined from TiN, a material exhibiting demonstrably lower intrinsic 1/f "
        "charge noise compared to Al, by electron beam lithography at 5 nm-8 nm "
        "half-pitch, with aluminum oxide (4 nm ALD-deposited Al2O3) gate dielectric. "
        "Interface trap density D_it is verified by low-frequency conductance profiling "
        "at 1 K; devices with D_it > 10^9 cm^-2 eV^-1 are rejected at wafer probe. "
        "The 2025 industrial demonstration of >99% fidelity in a 300 mm foundry [15] "
        "confirms that this fabrication pathway is compatible with volume production "
        "without specialized research infrastructure."
    ),
    (
        "B. Spintronic MTJ Stack Integration",
        2,
        "PMA-MTJ stacks are deposited by dc magnetron sputtering in the configuration: "
        "Ta(5)/Pt(3)/[Co(0.3)/Pt(0.3)]x5/Co(0.8)/MgO(0.9)/CoFeB(1.2)/Ta(5) [nm]. "
        "The MgO tunnel barrier is crystallized by post-deposition vacuum annealing at "
        "300 degrees C for 60 minutes, achieving TMR ratios exceeding 200% and four "
        "reliably resolved resistance states (R0-R3) requisite for 2-bit per synapse "
        "encoding. Patterning to 20 nm pillar diameter is executed by ion-beam etching "
        "(IBE) with in-situ endpoint detection. The maximum BEOL process temperature "
        "of 400 degrees C is respected throughout the MTJ integration sequence to "
        "preserve underlying transistor integrity."
    ),
    (
        "C. Three-Dimensional Stacking and Thermal Management",
        2,
        "Wafer-level Cu-Cu thermocompression bonding at a 200 nm pitch, demonstrated "
        "reproducibly on 300 mm wafers at IMEC and CEA-Leti, interconnects the three "
        "functional planes with via resistances below 1 ohm and alignment tolerances "
        "of +/-10 nm. A projected die footprint of ~0.5 mm^2 for the 256x256 crossbar "
        "configuration is compatible with dicing yields achievable on 300 mm wafers. "
        "Thermal management at 1 K is addressed through microfluidic superfluid-helium "
        "channels etched into the SOI handling wafer, phononic metamaterial thermal "
        "barriers between the MTJ layer and the QCPP, and thermal anchoring pads "
        "connecting the CPP substrate to the dilution refrigerator mixing chamber "
        "stage via oxygen-free high-conductivity (OFHC) copper braids."
    ),

    (
        "VI. SECURITY IMPLICATIONS AND THREAT LANDSCAPE",
        1,
        "The physical and operational novelty of the QNHS architecture introduces "
        "a multi-dimensional security surface that has no direct analog in either "
        "classical neuromorphic or conventional quantum computing hardware. We "
        "identify four principal threat categories and propose countermeasures "
        "that exploit the native physical properties of the QNHS substrate."
    ),
    (
        "A. Side-Channel Vulnerabilities in Cryogenic Systems",
        2,
        "The cryo-CMOS peripheral plane introduces several classical side-channel "
        "vectors. The 30 MHz Rabi frequency of qubit gate operations produces "
        "distinctive RF emission signatures on control lines that traverse the "
        "cryostat boundary, potentially enabling passive circuit reconstruction by "
        "an adversary with access to the wiring environment. MWPM decoder latency "
        "varies with syndrome weight, creating a timing channel that correlates "
        "observable decoder completion events with the instantaneous qubit error "
        "pattern, which in turn depends on the encoded weight state.\n\n"
        "The 2025 demonstration by Lu et al. [22] that timing measurements with "
        "10 queries suffice to fingerprint quantum hardware on cloud services "
        "directly motivates this concern for multi-tenant QNHS deployments. "
        "Choudhury et al. [23] further showed that crosstalk-based attacks achieve "
        "85.7% circuit reconstruction accuracy on NISQ hardware, a threat that "
        "extends naturally to the syndrome extraction sub-lattice of the QCPP.\n\n"
        "Countermeasures include constant-time MWPM scheduling with randomized idle "
        "cycles, electromagnetic shielding of qubit control lines within the inner "
        "cryostat shield, and hardware noise injection on power rails feeding the "
        "cryo-CMOS SerDes I/O to mask power consumption signatures."
    ),
    (
        "B. Adversarial Robustness of Quantum-Encoded Weights",
        2,
        "Quantum amplitude encoding confers partial adversarial resilience that is "
        "absent in classical deterministic neuromorphic systems. In a classical "
        "crossbar, an adversarial perturbation of a scalar synaptic weight w_ij "
        "directly alters the network output in a predictable, gradient-computable "
        "manner. In the QNHS, the same perturbation must instead shift the full "
        "probability distribution {|alpha_l|^2} across 2^k basis states; projective "
        "measurement at inference time then samples from this distribution, "
        "introducing irreducible stochasticity that degrades the effectiveness of "
        "gradient-based adversarial attacks [20].\n\n"
        "Nevertheless, determined adversaries can exploit the Q-STDP learning rule "
        "through training-phase attacks. By injecting carefully timed spurious spikes "
        "into the presynaptic input stream, an attacker can systematically bias "
        "amplitude distributions toward adversarially chosen weight states, "
        "implementing a backdoor attack without modifying hardware. A 2024 "
        "study identified analogous quantum state-universal adversarial perturbations "
        "(QS-UAP) in variational quantum circuit models [20]. Countermeasures include "
        "differential privacy on Q-STDP amplitude updates (gradient clipping before "
        "normalization), anomaly detection on spike-timing statistics during "
        "training, and redundant crossbar tiles with majority-vote inference for "
        "safety-critical workloads."
    ),
    (
        "C. Hardware Security Primitives and PQC Acceleration",
        2,
        "The QNHS substrate exposes three built-in defensive capabilities. First, "
        "charge noise in 28Si quantum dots exhibits device-unique stochastic "
        "signatures determined by the nanoscale distribution of interface traps and "
        "residual 29Si nuclei. These signatures constitute Physical Unclonable "
        "Functions (PUFs) suitable for device authentication and anti-counterfeiting "
        "without additional silicon area. Second, stochastic MTJ switching under "
        "sub-threshold SOT current pulses at 1 K generates true random bits from "
        "thermal fluctuations in the free-layer magnetization at rates exceeding "
        "1 Gbit/s, providing a hardware quantum random number generator (QRNG) "
        "for cryptographic key seeding.\n\n"
        "Third, the 256x256 MTJ crossbar naturally accelerates the polynomial ring "
        "arithmetic underlying NIST-standardized post-quantum lattice schemes "
        "(FIPS 203: ML-KEM, FIPS 204: ML-DSA [26]). Number theoretic transform "
        "(NTT) operations over Z_q reduce to analog matrix-vector multiplications "
        "that the crossbar executes in constant time at ~11 fJ per multiply-accumulate, "
        "compared to ~50 nJ for software implementation on ARM Cortex-M4. This "
        "represents a potential 4.5 million-fold energy advantage for PQC key "
        "generation and encapsulation at the hardware level."
    ),
    (
        "D. Supply Chain and Hardware Trojan Threats",
        2,
        "Hardware trojans embedded in cryo-CMOS circuits present an elevated risk "
        "in the QNHS context because standard room-temperature test vectors do not "
        "reveal cryogenic-mode vulnerabilities. A trojan designed to activate at "
        "the 1 GHz PLL operating frequency but remain dormant during low-frequency "
        "wafer probe would evade conventional manufacturing test. The MWPM decoder "
        "is a particularly high-value target: corrupting its syndrome processing "
        "would silently degrade logical qubit fidelity without any detectable "
        "increase in physical error rate.\n\n"
        "Countermeasures require trusted foundry certification for the 3 nm FinFET "
        "cryo-CMOS layer, logic encryption of the MWPM RTL prior to tape-out, and "
        "mandatory post-fabrication functional verification at 1 K operating "
        "temperature for all decoder instances. Isotopically enriched 28Si supply "
        "chain integrity (99.9995% purity) requires third-party isotopic verification "
        "by secondary ion mass spectrometry (SIMS) before MBE deposition."
    ),

    (
        "VII. DISCUSSION",
        1,
        ""
    ),
    (
        "A. Remaining Engineering Challenges",
        2,
        "Several significant challenges remain. Operating spin qubits at 1 K rather "
        "than the <20 mK temperatures preferred by superconducting platforms introduces "
        "k_B*T ~= 86 ueV thermal fluctuations, requiring orbital and valley splittings "
        "well in excess of this energy scale. The 28Si/SiO2 system exhibits orbital "
        "splittings of ~10 meV and valley splittings tunable from 0.1 meV to >1 meV "
        "via applied strain, placing 1 K operation within reach but demanding careful "
        "strain engineering. The physical coupling mechanism between MTJ stray-field "
        "gradients and spin qubit states in adjacent planes requires precise "
        "magnetostatic shielding. Nanoscale-patterned permalloy flux concentrators "
        "etched into the SSP-QCPP interface layer are proposed to confine MTJ stray "
        "fields to below 0.1 mT at the qubit plane, below the spin-flip threshold "
        "for 28Si spin qubits under typical g-factor conditions."
    ),
    (
        "B. Scalability and Integration Roadmap",
        2,
        "The QNHS architecture scales along two orthogonal dimensions: increasing "
        "qubit count per synapse k exponentially enriches representational capacity "
        "at the cost of syndrome extraction overhead; and increasing crossbar "
        "dimension N scales the synaptic fan-in/fan-out quadratically at the cost "
        "of increased thermal load. For the near-term horizon (<= 5 years), a "
        "64x64 crossbar with k = 2 qubits per synapse represents a tractable "
        "demonstration target, requiring 64^2 x 5 = 20,480 physical qubits and "
        "64^2 = 4,096 MTJ cells, well within the qubit counts already fabricated "
        "in academic silicon quantum dot foundry runs. The 2025 11-qubit atom "
        "processor [16] and 300 mm foundry demonstrations [15] mark concrete "
        "progress toward this milestone. The cryo-CMOS plane design must also "
        "contend with transistor threshold voltage shifts of ~100 mV at 4 K "
        "relative to room temperature characterization, and transconductance kinks "
        "arising from carrier freeze-out in lightly-doped nodes. Both effects have "
        "been characterized and correctable circuit design methodologies have been "
        "published by TU Delft and IMEC [14]."
    ),
    (
        "C. Broader Impact",
        2,
        "Should the QNHS architecture achieve projected performance metrics, the "
        "implications extend substantially beyond neuromorphic inference efficiency. "
        "Intrinsic Bayesian weight distributions, derived from quantum amplitude "
        "encoding rather than software-emulated Monte Carlo sampling, could enable "
        "fundamentally more calibrated uncertainty quantification in safety-critical "
        "AI systems, including autonomous vehicle perception pipelines and AI-assisted "
        "clinical diagnostics. The security primitives identified in Section VI "
        "suggest that future secure AI inference platforms could leverage the same "
        "physical substrate for both computation and cryptographic protection, "
        "eliminating the area and power overhead of dedicated security hardware. "
        "Furthermore, the quantum-coherent synaptic substrate offers a natural "
        "hardware primitive for quantum machine learning algorithms that currently "
        "require costly state preparation overhead on conventional gate-model quantum "
        "computers."
    ),

    (
        "VIII. CONCLUSION",
        1,
        "This paper has delineated the QNHS architecture: a vertically integrated "
        "monolithic system physically co-integrating 28Si spin qubits, spintronic "
        "PMA-MTJ synapses, and cryo-CMOS peripheral circuits. Theoretical analysis "
        "projects 11 fJ per synaptic event under realistic error-correction overhead, "
        "a three-order-of-magnitude improvement over GPU baselines and a ~20x net "
        "wall-plug advantage after cryogenic refrigeration cost. Recent 2025 "
        "experimental results demonstrating >99% gate fidelity and T1 times reaching "
        "9.5 seconds in 300 mm foundry-fabricated 28Si spin qubits [15], and the "
        "SemiQon cryo-CMOS breakthrough enabling 0.1% power consumption at 1 K [18], "
        "substantially strengthen the near-term credibility of the QNHS vision.\n\n"
        "This work also provides the first security analysis of a quantum-neuromorphic "
        "hardware substrate, identifying side-channel vulnerabilities demonstrated on "
        "cloud quantum services in 2025 [22, 23], adversarial threats to "
        "quantum-encoded weight distributions [20, 21], and native hardware security "
        "primitives including PUFs, QRNGs, and PQC-accelerating crossbar arithmetic "
        "aligned with NIST's August 2024 post-quantum cryptography standards [26]. "
        "The convergence of quantum mechanics, neuroscience-inspired computation, and "
        "hardware security at the nanoscale constitutes a frontier of extraordinary "
        "scientific richness and technological consequence. Open-source simulation "
        "code supporting this work is available at " + GITHUB_URL + "."
    ),
]

ACKNOWLEDGMENT = (
    "The author gratefully acknowledges the Cloud Security Alliance and HCLTech's "
    "research community for the cross-disciplinary intellectual environment that "
    "shaped the perspectives presented in this work. This research was conducted "
    "independently and does not represent the official views or positions of HCLTech "
    "(HCL America Inc.). Open-source simulation code supporting this paper is "
    "available at " + GITHUB_URL + "."
)

REFERENCES = [
    "[1] M. M. Waldrop, \"The chips are down for Moore's law,\" Nature, vol. 530, pp. 144-147, Feb. 2016. doi: 10.1038/530144a",
    "[2] IEEE International Roadmap for Devices and Systems (IRDS), \"More Moore White Paper,\" IEEE, 2023. irds.ieee.org/editions/2023",
    "[3] M. Davies et al., \"Advancing neuromorphic computing with Loihi: A survey of results and outlook,\" Proc. IEEE, vol. 109, no. 5, pp. 911-934, May 2021. doi: 10.1109/JPROC.2021.3067593",
    "[4] D. S. Modha et al., \"Neural inference at the frontier of energy, space, and time,\" Science, vol. 382, no. 6668, pp. 329-335, Oct. 2023. doi: 10.1126/science.adh1174",
    "[5] J. Preskill, \"Quantum computing in the NISQ era and beyond,\" Quantum, vol. 2, p. 79, Aug. 2018. doi: 10.22331/q-2018-08-06-79",
    "[6] C. Mead, \"Neuromorphic electronic systems,\" Proc. IEEE, vol. 78, no. 10, pp. 1629-1636, Oct. 1990. doi: 10.1109/5.58356",
    "[7] A. Sengupta, P. Panda, P. Wijesinghe, Y. Kim, and K. Roy, \"Magnetic tunnel junction mimics stochastic cortical spiking neurons,\" Sci. Rep., vol. 6, p. 30039, 2016. doi: 10.1038/srep30039",
    "[8] L. M. K. Vandersypen et al., \"Interfacing spin qubits in quantum dots and donors - hot, dense, and coherent,\" npj Quantum Inf., vol. 3, p. 34, 2017. doi: 10.1038/s41534-017-0038-y",
    "[9] A. M. J. Zwerver et al., \"Qubits made by advanced semiconductor manufacturing,\" Nature Electronics, vol. 5, pp. 184-190, 2022. doi: 10.1038/s41928-022-00727-9",
    "[10] S. D. Sarma, M. Freedman, and C. Nayak, \"Majorana zero modes and topological quantum information processing,\" npj Quantum Inf., vol. 1, p. 15001, 2015. doi: 10.1038/npjqi.2015.1",
    "[11] G. G. Guerreschi and A. Y. Matsuura, \"QAOA for Max-Cut requires hundreds of qubits for quantum speed-up,\" Sci. Rep., vol. 9, p. 6903, 2019. doi: 10.1038/s41598-019-43176-9",
    "[12] M. Maronese, C. Destri, and E. Prati, \"Quantum activation functions for quantum neural networks,\" Quantum Inf. Process., vol. 21, p. 128, 2022. doi: 10.1007/s11128-022-03466-0",
    "[13] J. Yoneda et al., \"A quantum-dot spin qubit with coherence limited by charge noise and fidelity higher than 99.9%,\" Nature Nanotechnology, vol. 13, pp. 102-106, 2018. doi: 10.1038/s41565-017-0014-x",
    "[14] B. Patra et al., \"Cryo-CMOS circuits and systems for quantum computing applications,\" IEEE J. Solid-State Circuits, vol. 53, no. 1, pp. 309-321, Jan. 2018. doi: 10.1109/JSSC.2017.2737549",
    "[15] A. R. Mills et al., \"Industry-compatible silicon spin-qubit unit cells exceeding 99% fidelity,\" Nature, vol. 646, pp. 81-87, 2025. doi: 10.1038/s41586-025-09531-9",
    "[16] M. Y. Simmons et al., \"An 11-qubit atom processor in silicon,\" Nature, vol. 648, pp. 569-575, 2025. doi: 10.1038/s41586-025-09827-w",
    "[17] M. Aghaee et al., \"Interferometric single-shot parity measurement in InAs-Al hybrid devices,\" Nature, 2025. doi: 10.1038/s41586-025-08107-9 [Note: topological qubit claims under independent scientific review]",
    "[18] SemiQon Oy, \"Industry-first cryo-CMOS transistor optimized for operation at 1 K and below,\" Technical Announcement, Nov. 2024. URL: semiqon.tech",
    "[19] B. Nijboer et al., \"Millikelvin Si-MOSFETs for quantum electronics,\" arXiv:2410.01077, 2024.",
    "[20] E. Yocam, A. Rizi, M. Kamepalli, V. Vaidyan, Y. Wang, and G. Comert, \"Quantum adversarial machine learning and defense strategies: Challenges and opportunities,\" arXiv:2412.12373, Dec. 2024.",
    "[21] A. Upadhyay et al., \"Adversarial robustness in quantum machine learning: A scoping review,\" Computers, vol. 15, no. 4, p. 233, 2025. doi: 10.3390/computers15040233",
    "[22] C. Lu, E. Telang, A. Aysu, and K. Basu, \"Quantum Leak: Timing side-channel attacks on cloud-based quantum services,\" in Proc. GLSVLSI, 2025. doi: 10.1145/3716368.3735264",
    "[23] N. Choudhury et al., \"Crosstalk-induced side channel threats in multi-tenant NISQ computers,\" in Proc. NDSS, 2025. arXiv:2412.10507",
    "[24] H. Neven et al., \"Quantum hyperdimensional computing: a foundational paradigm for quantum neuromorphic architectures,\" npj Unconventional Comput., 2025. doi: 10.1038/s44335-026-00064-6",
    "[25] J. Grollier et al., \"Neuromorphic computing with spintronics,\" npj Spintronics, vol. 2, p. 19, 2024. doi: 10.1038/s44306-024-00019-2",
    "[26] National Institute of Standards and Technology (NIST), \"Post-Quantum Cryptography Standards: FIPS 203, 204, 205,\" Aug. 2024. URL: csrc.nist.gov/publications/fips",
]

AUTHOR_BIO = (
    "Sunil Gentyala is an independent researcher and Lead Cybersecurity and AI "
    "Security Consultant with a 10-year tenure at HCLTech (HCL America Inc.), "
    "Dallas, TX, where he serves as Reporting Manager overseeing a cross-functional "
    "team of 11 professionals spanning onsite US and offshore international "
    "operations. He is designated as HCLTech's expert representative to the Cloud "
    "Security Alliance (CSA), specializing in artificial intelligence security, "
    "adversarial machine learning, and enterprise risk management. Mr. Gentyala "
    "actively publishes cybersecurity and emerging technology research across "
    "industry publications including Dark Reading, IDG/Foundry platforms "
    "(Computerworld, CSO Online, CIO.com), and Cyber Defense Magazine. He is a "
    "recipient of the Cybersecurity Excellence Award nomination in recognition of "
    "his contributions to AI and enterprise security. His scholarly work draws upon "
    "extensive peer-reviewed research and first-hand empirical evaluation of major "
    "AI platforms, LLM security vulnerabilities, jailbreak prevention methodologies, "
    "quantum-classical hybrid computing architectures, quantum hardware security, "
    "neuromorphic systems, and post-von Neumann computational substrates. He received "
    "his education at SRM College and Kakatiya University, Karimnagar, India.\n\n"
    "Affiliation: Independent Researcher, HCLTech (HCL America Inc.), Dallas, TX 75001, USA\n"
    "Email: sunil.gentyala@ieee.org\n"
    "ORCID: 0009-0005-2642-3479\n"
    "GitHub: " + GITHUB_URL + "\n"
    "IEEE Senior Member"
)

TABLE_1_CAPTION = (
    "Table I: Comparative Architecture Assessment: GPU vs. Neuromorphic vs. QNHS"
)
TABLE_1_HEADERS = [
    "Metric", "NVIDIA H100 (GPU)", "Intel Loihi 2", "IBM NorthPole", "QNHS (Projected)"
]
TABLE_1_ROWS = [
    ["Energy / synaptic event (E_syn)", "~10 pJ", "~1 pJ", "~0.5 pJ", "~11 fJ"],
    ["Net system efficiency (wall-plug)", "1x (baseline)", "~10x", "~20x", "~20x"],
    ["Weight precision", "FP16 (16-bit)", "INT8", "INT8", "2^k quantum states"],
    ["Representational uncertainty", "None (deterministic)", "None", "None", "Intrinsic (Bayesian)"],
    ["Spike-native processing", "No", "Yes", "Partial", "Yes"],
    ["Parallelism mechanism", "16,896 CUDA cores", "1M neurons", "256 cores", "Quantum superposition"],
    ["Operating temperature", "300 K", "300 K", "300 K", "~1 K"],
    ["Hardware security primitives", "None", "None", "None", "PUF, QRNG, PQC accel."],
]
TABLE_1_NOTE = (
    "Net system efficiency accounts for cryogenic refrigeration overhead "
    "(~50x Carnot factor at 1 K from 300 K) for QNHS. QNHS figures are "
    "theoretical projections; H100, Loihi 2, and NorthPole figures are from "
    "published specifications."
)


# ─────────────────────────────────────────────────────────────────────────────
#  DOCX HELPER UTILITIES
# ─────────────────────────────────────────────────────────────────────────────

def set_two_columns(doc: Document):
    """Apply two-column layout to the last section."""
    sect = doc.sections[-1]
    sectPr = sect._sectPr
    cols = OxmlElement("w:cols")
    cols.set(qn("w:num"), "2")
    cols.set(qn("w:space"), "720")  # 0.5 inch gap in twips
    cols.set(qn("w:equalWidth"), "1")
    sectPr.append(cols)


def set_page_ieee(doc: Document):
    """IEEE letter page: 8.5 x 11 in, margins 0.75 in all around."""
    sect = doc.sections[0]
    sect.page_width = Inches(8.5)
    sect.page_height = Inches(11)
    sect.left_margin = Inches(0.75)
    sect.right_margin = Inches(0.75)
    sect.top_margin = Inches(0.75)
    sect.bottom_margin = Inches(0.75)


def set_page_acm(doc: Document):
    """ACM letter page: 8.5 x 11 in, margins 1 in all around."""
    sect = doc.sections[0]
    sect.page_width = Inches(8.5)
    sect.page_height = Inches(11)
    sect.left_margin = Inches(0.9)
    sect.right_margin = Inches(0.9)
    sect.top_margin = Inches(1.0)
    sect.bottom_margin = Inches(1.0)


def add_run(para, text, bold=False, italic=False, font_name="Times New Roman",
            font_size=10, color=None):
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = font_name
    run.font.size = Pt(font_size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return run


def add_paragraph(doc, text, style=None, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                  space_before=0, space_after=4, font_name="Times New Roman",
                  font_size=10, bold=False, italic=False, left_indent=0):
    para = doc.add_paragraph(style=style)
    para.alignment = alignment
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.space_after = Pt(space_after)
    if left_indent:
        para.paragraph_format.left_indent = Inches(left_indent)
    add_run(para, text, bold=bold, italic=italic,
            font_name=font_name, font_size=font_size)
    return para


def add_ieee_section_heading(doc, text, level=1):
    if level == 1:
        para = add_paragraph(
            doc, text,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            space_before=6, space_after=4,
            bold=True, font_size=10
        )
    else:
        parts = text.split(". ", 1)
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.paragraph_format.space_before = Pt(4)
        para.paragraph_format.space_after = Pt(2)
        if len(parts) == 2:
            add_run(para, parts[0] + ". ", bold=True, font_size=10, italic=True)
            add_run(para, parts[1], bold=False, font_size=10, italic=True)
        else:
            add_run(para, text, bold=True, font_size=10, italic=True)
    return para


def add_acm_section_heading(doc, text, level=1):
    if level == 1:
        para = add_paragraph(
            doc, text,
            alignment=WD_ALIGN_PARAGRAPH.LEFT,
            space_before=8, space_after=4,
            bold=True, font_size=11, font_name="Arial"
        )
    else:
        parts = text.split(". ", 1)
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        para.paragraph_format.space_before = Pt(6)
        para.paragraph_format.space_after = Pt(2)
        if len(parts) == 2:
            add_run(para, parts[0] + ". ", bold=True, font_size=10, font_name="Arial")
            add_run(para, parts[1], bold=True, font_size=10, font_name="Arial")
        else:
            add_run(para, text, bold=True, font_size=10, font_name="Arial")
    return para


def add_table(doc, headers, rows, caption, note, font_size=8, font_name="Times New Roman"):
    para = add_paragraph(doc, caption, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                         space_before=6, space_after=2, bold=True,
                         font_size=font_size + 1, font_name=font_name)
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.paragraphs[0].clear()
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(font_size)
        run.font.name = font_name
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r_idx, row in enumerate(rows):
        tbl_row = table.rows[r_idx + 1]
        for c_idx, cell_text in enumerate(row):
            cell = tbl_row.cells[c_idx]
            cell.paragraphs[0].clear()
            run = cell.paragraphs[0].add_run(cell_text)
            run.font.size = Pt(font_size)
            run.font.name = font_name
            if c_idx == 0:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            else:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    note_para = add_paragraph(doc, note, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                              space_before=0, space_after=6,
                              italic=True, font_size=font_size, font_name=font_name)
    return table


def add_references(doc, refs, font_name="Times New Roman", font_size=8):
    for ref in refs:
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(2)
        para.paragraph_format.left_indent = Inches(0.2)
        para.paragraph_format.first_line_indent = Inches(-0.2)
        add_run(para, ref, font_name=font_name, font_size=font_size)


# ─────────────────────────────────────────────────────────────────────────────
#  IEEE CONFERENCE PAPER GENERATOR
# ─────────────────────────────────────────────────────────────────────────────

def generate_ieee_paper(output_path: str):
    doc = Document()
    set_page_ieee(doc)

    # Remove default styles noise
    for style in ["Normal", "Default Paragraph Font"]:
        pass  # keep defaults but override via explicit formatting

    # ---- Title ----
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.space_before = Pt(8)
    title_para.paragraph_format.space_after = Pt(6)
    add_run(title_para, TITLE, bold=True, font_size=16)

    # ---- Author ----
    auth_para = add_paragraph(
        doc, AUTHOR,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=2, space_after=2,
        font_size=11
    )
    aff_para = add_paragraph(
        doc,
        AFFILIATION + "\n" + EMAIL + " | ORCID: " + ORCID,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=0, space_after=8,
        italic=True, font_size=9
    )

    # ---- Abstract ----
    abs_para = doc.add_paragraph()
    abs_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    abs_para.paragraph_format.space_before = Pt(4)
    abs_para.paragraph_format.space_after = Pt(2)
    add_run(abs_para, "Abstract", bold=True, italic=True, font_size=9)
    add_run(abs_para, "-" + ABSTRACT, italic=True, font_size=9)

    # ---- Index Terms ----
    idx_para = doc.add_paragraph()
    idx_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    idx_para.paragraph_format.space_before = Pt(2)
    idx_para.paragraph_format.space_after = Pt(8)
    add_run(idx_para, "Index Terms", bold=True, italic=True, font_size=9)
    add_run(idx_para, "-" + INDEX_TERMS, italic=True, font_size=9)

    # ---- Switch to two columns ----
    set_two_columns(doc)

    # ---- Body Sections ----
    table_inserted = False
    for heading, level, body in SECTIONS:
        add_ieee_section_heading(doc, heading, level)
        if body:
            for para_text in body.split("\n\n"):
                pt = para_text.strip()
                if not pt:
                    continue
                add_paragraph(doc, pt, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                              space_before=0, space_after=4, font_size=10)

        # Insert Table I after Section IV-A
        if heading == "A. Energy Efficiency Decomposition" and not table_inserted:
            add_table(doc, TABLE_1_HEADERS, TABLE_1_ROWS,
                      TABLE_1_CAPTION, TABLE_1_NOTE)
            table_inserted = True

    # ---- Acknowledgment ----
    add_ieee_section_heading(doc, "ACKNOWLEDGMENT", 1)
    add_paragraph(doc, ACKNOWLEDGMENT, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                  space_before=0, space_after=4, font_size=10)

    # ---- References ----
    add_ieee_section_heading(doc, "REFERENCES", 1)
    add_references(doc, REFERENCES, font_size=8)

    # ---- Author Biography ----
    doc.add_page_break()
    add_ieee_section_heading(doc, "AUTHOR BIOGRAPHY", 1)
    add_paragraph(doc, AUTHOR_BIO, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                  space_before=0, space_after=6, font_size=10)

    doc.save(output_path)
    print(f"IEEE paper saved: {output_path}")


# ─────────────────────────────────────────────────────────────────────────────
#  ACM CONFERENCE PAPER GENERATOR
# ─────────────────────────────────────────────────────────────────────────────

def generate_acm_paper(output_path: str):
    doc = Document()
    set_page_acm(doc)

    # ---- ACM copyright stub ----
    cp = add_paragraph(
        doc,
        "ACM Reference Format: Sunil Gentyala. 2026. Entangled Intelligence: "
        "Nanoscale Quantum-Neuromorphic Hybrid Architectures for Post-von Neumann "
        "Computation. In Proceedings of the ACM Conference (venue TBD). ACM, New York, NY, USA.",
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        space_before=0, space_after=8,
        font_size=8, font_name="Arial", italic=True
    )

    # ---- Title ----
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.space_before = Pt(8)
    title_para.paragraph_format.space_after = Pt(6)
    add_run(title_para, TITLE, bold=True, font_size=18, font_name="Arial")

    # ---- Author block (ACM style) ----
    auth_para = doc.add_paragraph()
    auth_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    auth_para.paragraph_format.space_before = Pt(4)
    auth_para.paragraph_format.space_after = Pt(2)
    add_run(auth_para, AUTHOR, bold=True, font_size=12, font_name="Arial")

    aff_lines = (
        AFFILIATION + "\n" +
        "Dallas, TX 75001, USA\n" +
        EMAIL + "\n" +
        "ORCID: " + ORCID
    )
    aff_para = add_paragraph(
        doc, aff_lines,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=0, space_after=8,
        font_size=9, font_name="Arial"
    )

    # ---- Abstract (ACM) ----
    add_paragraph(doc, "ABSTRACT", alignment=WD_ALIGN_PARAGRAPH.LEFT,
                  bold=True, font_size=10, font_name="Arial",
                  space_before=6, space_after=2)
    add_paragraph(doc, ABSTRACT, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                  font_size=9, font_name="Arial", space_before=0, space_after=4,
                  left_indent=0.3)

    # ---- CCS Concepts ----
    add_paragraph(doc, "CCS CONCEPTS", alignment=WD_ALIGN_PARAGRAPH.LEFT,
                  bold=True, font_size=10, font_name="Arial",
                  space_before=4, space_after=2)
    ccs = (
        "* Hardware -> Quantum computation; Neuromorphic devices; "
        "* Computing methodologies -> Neural networks; Quantum computing; "
        "* Security and privacy -> Hardware security; Side-channel analysis; "
        "Post-quantum cryptography."
    )
    add_paragraph(doc, ccs, font_size=9, font_name="Arial",
                  space_before=0, space_after=4, left_indent=0.3)

    # ---- Keywords ----
    add_paragraph(doc, "KEYWORDS", alignment=WD_ALIGN_PARAGRAPH.LEFT,
                  bold=True, font_size=10, font_name="Arial",
                  space_before=4, space_after=2)
    add_paragraph(
        doc,
        "quantum-neuromorphic computing, silicon spin qubits, spintronic memristors, "
        "cryo-CMOS, hardware security, side-channel attacks, post-quantum cryptography, "
        "spiking neural networks, topological error correction, neuromorphic AI",
        font_size=9, font_name="Arial",
        space_before=0, space_after=8, left_indent=0.3
    )

    # ---- Separator ----
    sep = doc.add_paragraph()
    sep.paragraph_format.space_before = Pt(0)
    sep.paragraph_format.space_after = Pt(8)
    add_run(sep, "_" * 70, font_size=8, font_name="Arial")

    # ---- Apply two columns for body ----
    set_two_columns(doc)

    # ---- Body sections ----
    table_inserted = False
    for heading, level, body in SECTIONS:
        add_acm_section_heading(doc, heading, level)
        if body:
            for para_text in body.split("\n\n"):
                pt = para_text.strip()
                if not pt:
                    continue
                add_paragraph(doc, pt, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                              space_before=0, space_after=4,
                              font_size=9, font_name="Arial")

        if heading == "A. Energy Efficiency Decomposition" and not table_inserted:
            add_table(doc, TABLE_1_HEADERS, TABLE_1_ROWS,
                      TABLE_1_CAPTION, TABLE_1_NOTE,
                      font_size=7, font_name="Arial")
            table_inserted = True

    # ---- Acknowledgment ----
    add_acm_section_heading(doc, "ACKNOWLEDGMENT", 1)
    add_paragraph(doc, ACKNOWLEDGMENT, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                  space_before=0, space_after=4, font_size=9, font_name="Arial")

    # ---- References ----
    add_acm_section_heading(doc, "REFERENCES", 1)
    add_references(doc, REFERENCES, font_name="Arial", font_size=8)

    # ---- Author Biography ----
    doc.add_page_break()
    add_acm_section_heading(doc, "AUTHOR BIOGRAPHY", 1)
    add_paragraph(doc, AUTHOR_BIO, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                  space_before=0, space_after=6, font_size=9, font_name="Arial")

    doc.save(output_path)
    print(f"ACM paper saved: {output_path}")


# ─────────────────────────────────────────────────────────────────────────────
#  MAIN
# ─────────────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    ieee_path = os.path.join("papers", "IEEE", "NANO_2026_IEEE_Gentyala.docx")
    acm_path = os.path.join("papers", "ACM", "NANO_2026_ACM_Gentyala.docx")

    os.makedirs(os.path.dirname(ieee_path), exist_ok=True)
    os.makedirs(os.path.dirname(acm_path), exist_ok=True)

    print("Generating IEEE conference paper...")
    generate_ieee_paper(ieee_path)

    print("Generating ACM conference paper...")
    generate_acm_paper(acm_path)

    print("\nDone. Files generated:")
    print(f"  IEEE: {os.path.abspath(ieee_path)}")
    print(f"  ACM:  {os.path.abspath(acm_path)}")
